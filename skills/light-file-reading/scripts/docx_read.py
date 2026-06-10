#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""DOCX 读取与结构理解（light-file-reading）。

依赖：python-docx。读段落/标题/run 样式、表格、节页边距、内嵌图计数。
注意：python-docx 不读修订(tracked changes)、不渲染。需要修订/批注请走
references/DOCX-REF.md 的 pandoc --track-changes=all 或裸 XML 路线。
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")


def read_paragraphs(path):
    """返回 [(style_name, text), ...]，跳过空段。用于看章节骨架。"""
    from docx import Document
    doc = Document(path)
    return [(p.style.name, p.text) for p in doc.paragraphs if p.text.strip()]


def read_headings(path):
    """只取标题段 → [(level, text)]，level 从 style 名 'Heading N' 解析。"""
    out = []
    for style, text in read_paragraphs(path):
        if style and style.startswith("Heading"):
            try:
                lvl = int(style.split()[-1])
            except ValueError:
                lvl = 1
            out.append((lvl, text))
        elif style == "Title":
            out.append((0, text))
    return out


def read_runs(path):
    """逐段逐 run 读样式（bold/italic/size/font），用于提取格式要求。
    返回 [{'para': i, 'text':.., 'bold':.., 'italic':.., 'size':.., 'font':..}]。"""
    from docx import Document
    doc = Document(path)
    rows = []
    for i, p in enumerate(doc.paragraphs):
        for r in p.runs:
            if not r.text:
                continue
            sz = r.font.size
            rows.append({
                "para": i,
                "text": r.text,
                "bold": r.bold,
                "italic": r.italic,
                "size": sz.pt if sz is not None else None,
                "font": r.font.name,
            })
    return rows


def read_tables(path):
    """读所有表格 → [list[list[str]]]（含表头行）。"""
    from docx import Document
    doc = Document(path)
    out = []
    for t in doc.tables:
        out.append([[c.text for c in row.cells] for row in t.rows])
    return out


def read_layout(path):
    """读节级页面格式（页边距/纸张），用于提取模板规范。返回 [dict]。"""
    from docx import Document
    doc = Document(path)

    def emu_in(v):
        return round(v.inches, 3) if v is not None else None  # Length→inch

    secs = []
    for s in doc.sections:
        secs.append({
            "page_w_in": emu_in(s.page_width),
            "page_h_in": emu_in(s.page_height),
            "margin_top_in": emu_in(s.top_margin),
            "margin_bottom_in": emu_in(s.bottom_margin),
            "margin_left_in": emu_in(s.left_margin),
            "margin_right_in": emu_in(s.right_margin),
        })
    return secs


def _selftest():
    """合成 docx 跑全流程，断言后清理。"""
    import os
    import tempfile
    from docx import Document
    from docx.shared import Pt

    tmp = tempfile.mkdtemp(prefix="docxread_")
    src = os.path.join(tmp, "sample.docx")
    doc = Document()
    doc.add_heading("Main Title", 0)
    doc.add_heading("Section One", 1)
    p = doc.add_paragraph()
    r = p.add_run("bold words")
    r.bold = True
    r.font.size = Pt(14)
    doc.add_heading("Section Two", 2)
    doc.add_paragraph("plain body text")
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "H1"
    t.rows[0].cells[1].text = "H2"
    t.rows[1].cells[0].text = "a"
    t.rows[1].cells[1].text = "b"
    doc.save(src)

    paras = read_paragraphs(src)
    assert any("plain body text" == txt for _, txt in paras), paras

    heads = read_headings(src)
    assert (0, "Main Title") in heads and (1, "Section One") in heads, heads

    runs = read_runs(src)
    assert any(rr["bold"] and rr["size"] == 14.0 for rr in runs), runs

    tables = read_tables(src)
    assert tables and tables[0][0] == ["H1", "H2"], tables

    layout = read_layout(src)
    assert layout and layout[0]["margin_top_in"] is not None, layout

    import shutil
    shutil.rmtree(tmp, ignore_errors=True)
    print("docx_read self-test OK: paragraphs/headings/runs/tables/layout all passed")


if __name__ == "__main__":
    if len(sys.argv) > 2 or (len(sys.argv) == 2 and sys.argv[1] != "--selftest"):
        raise SystemExit(f"usage: python {sys.argv[0]} [--selftest]")
    _selftest()
